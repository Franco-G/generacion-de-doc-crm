import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# --- CONFIGURACIÓN ---
ARCHIVO_EXCEL = 'CRM - REQUERIMIENTOS.xlsx'
HOJA_DATOS = 'Requerimiento'
CARPETA_SALIDA = 'Constancias_Finales_CRM'

def limpiar_nombre_archivo(texto):
    """Limpia caracteres inválidos para nombres de archivo."""
    if pd.isna(texto): return "Sin_Nombre"
    texto = str(texto).replace('/', '-').replace('\\', '-')
    return re.sub(r'[^\w\s-]', '', texto).strip().replace(' ', '_')

def agregar_campo_negrita(doc, etiqueta, valor):
    """Agrega una línea con formato 'Etiqueta: Valor'."""
    p = doc.add_paragraph()
    run = p.add_run(etiqueta)
    run.bold = True
    p.add_run(f" {valor}")

def insertar_texto_multilinea(doc, texto):
    """
    Inserta texto respetando los saltos de línea del Excel.
    Si la celda está vacía o dice 'NO HAY', maneja el formato adecuado.
    """
    if pd.isna(texto):
        doc.add_paragraph("(Sin información detallada)")
        return

    texto_str = str(texto).strip()
    
    # Manejo específico si en el Excel pusiste "NO HAY"
    if texto_str.upper() == "NO HAY":
        doc.add_paragraph("(No hay observaciones actuales)")
        return
        
    if not texto_str:
        doc.add_paragraph("(Sin información detallada)")
        return

    # Dividir por saltos de línea para crear párrafos reales en Word
    parrafos = texto_str.split('\n')
    for p in parrafos:
        if p.strip():
            doc.add_paragraph(p.strip())

def generar_constancias():
    # 1. Crear carpeta
    if not os.path.exists(CARPETA_SALIDA):
        os.makedirs(CARPETA_SALIDA)

    print(f"Leyendo '{ARCHIVO_EXCEL}'...")
    
    try:
        df = pd.read_excel(ARCHIVO_EXCEL, sheet_name=HOJA_DATOS)
    except Exception as e:
        print(f"Error crítico al leer el Excel: {e}")
        return

    # 2. Filtrar filas vacías (Evita generar 90 archivos, solo los reales)
    # Filtramos filas que no tengan MÓDULO o SUB-MÓDULO
    df_clean = df.dropna(subset=['MÓDULO', 'SUB-MÓDULO']).copy()
    df_clean.reset_index(drop=True, inplace=True)
    
    print(f"Generando {len(df_clean)} documentos...")

    # 3. Iterar y generar
    for index, row in df_clean.iterrows():
        try:
            # --- DATOS BASE ---
            # Generar ID (REQ01, REQ02...)
            req_num = str(index + 1).zfill(2)
            req_id = f"REQ{req_num}"
            
            modulo = str(row.get('MÓDULO', '')).strip()
            submodulo = str(row.get('SUB-MÓDULO', '')).strip()

            # --- OBTENCIÓN DE TEXTOS LARGOS (Con nombres corregidos) ---
            # Usamos .get() por seguridad, pero buscamos la columna exacta
            evidencia = row.get('EVIDENCIA DE IMPLEMENTACION', '') 
            situacion = row.get('SITUACION ACTUAL / PROBLEMA IDENTIFICADO', '')
            alcance = row.get('ALCANCE FUNCIONAL', '')
            # Nota: Incluyo los dos puntos al final porque así aparecía en tu CSV anterior
            visualizacion = row.get('COMO DEBERIA VISUALIZARSE CON LOS CAMBIOS REQUERIDOS:', '')

            # --- CREAR DOCUMENTO ---
            doc = Document()
            
            # Estilos
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Encabezado
            t = doc.add_heading(f"Constancia de Implementación {req_num}", 0)
            t.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Metadatos
            agregar_campo_negrita(doc, "Titulo:", submodulo)
            doc.add_paragraph() # Espacio
            agregar_campo_negrita(doc, "IF/IC:", "Requerimiento")
            agregar_campo_negrita(doc, "Módulo:", f"{modulo} / {submodulo}")
            agregar_campo_negrita(doc, "Url:", "")
            
            doc.add_paragraph("\nRealizar las siguientes modificaciones:")

            # SECCIÓN 1: EVIDENCIA
            doc.add_heading('EVIDENCIA DE IMPLEMENTACIÓN', level=1)
            insertar_texto_multilinea(doc, evidencia)

            # SECCIÓN 2: SITUACIÓN ACTUAL
            doc.add_heading('Situación actual / Problema identificado', level=2)
            insertar_texto_multilinea(doc, situacion)

            # SECCIÓN 3: ALCANCE FUNCIONAL
            doc.add_heading('Alcance funcional', level=2)
            insertar_texto_multilinea(doc, alcance)

            # SECCIÓN 4: VISUALIZACIÓN
            doc.add_paragraph()
            p_vis = doc.add_paragraph("Cómo debería visualizarse con los cambios requeridos:")
            p_vis.runs[0].bold = True
            
            # Verificar si hay texto específico o placeholder
            if pd.isna(visualizacion) or str(visualizacion).strip() == "" or "(INSERTE" in str(visualizacion):
                 doc.add_paragraph("[Espacio reservado para imagen o mockup]")
            else:
                 insertar_texto_multilinea(doc, visualizacion)

            # Guardar archivo
            safe_sub = limpiar_nombre_archivo(submodulo)
            filename = f"{req_id}_{safe_sub}.docx"
            ruta_final = os.path.join(CARPETA_SALIDA, filename)
            
            doc.save(ruta_final)
            print(f"Generado: {filename}")

        except Exception as e:
            print(f"Error en fila {index} ({submodulo}): {e}")

    print(f"\n--- PROCESO FINALIZADO ---")
    print(f"Documentos guardados en: {CARPETA_SALIDA}")

if __name__ == "__main__":
    generar_constancias()